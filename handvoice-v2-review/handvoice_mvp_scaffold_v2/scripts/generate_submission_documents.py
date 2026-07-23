"""Generate the controlled one-page Nokia proposal draft as DOCX and PDF.

The source remains Markdown so unsupported claims and open gates stay visible in
version control. The PDF is a review artifact; the final submission must still
pass the font/render checklist on the submission machine.
"""
from __future__ import annotations

from html import escape
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Mm, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION_DIR = ROOT / "docs" / "submissions" / "nokia-stage2-2026"
SOURCE = SUBMISSION_DIR / "proposal-one-page-draft.md"
DOCX_OUT = SUBMISSION_DIR / "Nokia_Stage2_HandVoice_Proposal_DRAFT.docx"
PDF_OUT = SUBMISSION_DIR / "Nokia_Stage2_HandVoice_Proposal_DRAFT.pdf"
BUILD_REPORT = SUBMISSION_DIR / "document-build-report.txt"


def parse_markdown() -> tuple[str, list[tuple[str, str]]]:
    title = ""
    sections: list[tuple[str, str]] = []
    heading: str | None = None
    body: list[str] = []
    for raw in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if line.startswith("# "):
            title = line[2:].strip()
            continue
        if line.startswith("## "):
            if heading is not None:
                sections.append((heading, " ".join(body).strip()))
            heading = line[3:].strip()
            body = []
            continue
        if line:
            body.append(line)
    if heading is not None:
        sections.append((heading, " ".join(body).strip()))
    if not title or not sections:
        raise RuntimeError("proposal Markdown is missing a title or sections")
    return title, sections


def plain(text: str) -> str:
    return re.sub(r"[*`]", "", text)


def generate_docx(title: str, sections: list[tuple[str, str]]) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.42)
    section.right_margin = Inches(0.42)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(12)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run("DRAFT — NOT SUBMISSION-READY: evidence and budget gates remain open.")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(155, 28, 28)

    for heading, body in sections:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.space_before = Pt(1.5)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(heading)
        run.bold = True
        run.font.name = "Aptos"
        run.font.size = Pt(12)

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(plain(body))
        run.font.name = "Aptos"
        run.font.size = Pt(12)

    document.core_properties.title = title
    document.core_properties.subject = "Nokia Stage 2 controlled draft — not submission-ready"
    document.save(DOCX_OUT)


def generate_pdf(title: str, sections: list[tuple[str, str]]) -> int:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ProposalTitle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=12.4,
        spaceAfter=2,
        alignment=TA_LEFT,
    )
    heading_style = ParagraphStyle(
        "ProposalHeading",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=12.2,
        spaceBefore=1.5,
        spaceAfter=0,
        keepWithNext=True,
    )
    body_style = ParagraphStyle(
        "ProposalBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=12,
        leading=12.2,
        spaceBefore=0,
        spaceAfter=0,
        alignment=TA_LEFT,
    )
    draft_style = ParagraphStyle(
        "DraftNotice",
        parent=body_style,
        fontName="Helvetica-Bold",
        textColor="#9b1c1c",
    )
    document = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        leftMargin=0.35 * inch,
        rightMargin=0.35 * inch,
        topMargin=0.30 * inch,
        bottomMargin=0.30 * inch,
        title=title,
        subject="Controlled draft — not submission-ready",
    )
    story = [
        Paragraph(escape(title), title_style),
        Paragraph(
            "DRAFT — NOT SUBMISSION-READY: evidence and budget gates remain open.",
            draft_style,
        ),
    ]
    for heading, body in sections:
        story.append(Paragraph(escape(heading), heading_style))
        chosen_style = draft_style if heading.startswith("DRAFT") else body_style
        story.append(Paragraph(escape(plain(body)).replace("\n", " "), chosen_style))
    document.build(story)
    return len(PdfReader(str(PDF_OUT)).pages)


def main() -> None:
    title, sections = parse_markdown()
    proposal_text = " ".join([title, *(f"{heading} {body}" for heading, body in sections)])
    word_count = len(re.findall(r"\b[\w'-]+\b", plain(proposal_text)))
    generate_docx(title, sections)
    page_count = generate_pdf(title, sections)
    status = "PASS" if page_count == 1 else "FAIL"
    BUILD_REPORT.write_text(
        "\n".join(
            [
                f"PDF page-count check: {status} ({page_count} page(s))",
                f"Proposal word-count check: {'PASS' if 550 <= word_count <= 650 else 'FAIL'} ({word_count} words; target 550-650).",
                "DOCX requested font: Aptos, 12 pt, single-spaced, A4.",
                "PDF review font: Helvetica, 12 pt because Aptos is unavailable in the build runtime.",
                "FINAL GATE: export the DOCX to PDF on a machine with Aptos and confirm exactly one page.",
                "CONTENT GATE: draft remains non-submittable until OPEN evidence fields are resolved.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    if page_count != 1:
        raise SystemExit(f"generated PDF has {page_count} pages; shorten content, never shrink the font")
    if not 550 <= word_count <= 650:
        raise SystemExit(f"proposal has {word_count} words; target 550-650")
    print(DOCX_OUT)
    print(PDF_OUT)


if __name__ == "__main__":
    main()
